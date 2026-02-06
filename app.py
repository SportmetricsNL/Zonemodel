import os
import time
import streamlit as st
import google.generativeai as genai
import pypdf
import docx

# --- 0. PAGE & THEME -------------------------------------------------------
st.set_page_config(
    page_title="Sportfysioloog AI",
    page_icon="🚴‍♂️",
    layout="wide",
)

# Global look & feel
st.markdown(
    r"""
    <style>
      :root {
        --bg: linear-gradient(135deg, #0b1229 0%, #0e1f40 50%, #0b5c6f 100%);
        --card: rgba(255, 255, 255, 0.06);
        --glass: rgba(255, 255, 255, 0.08);
        --border: rgba(255, 255, 255, 0.15);
        --accent: #3ce37b;
        --accent-2: #42c5f5;
        --text: #e8f4ff;
        --muted: #9fb3d9;
      }
      .stApp { background: var(--bg); color: var(--text); }
      .block-container { padding: 1.5rem 2.5rem 3rem; max-width: 1100px; }
      .hero {
        background: var(--card);
        border: 1px solid var(--border);
        border-radius: 18px;
        padding: 1.25rem 1.5rem;
        box-shadow: 0 15px 50px rgba(0,0,0,0.35);
      }
      .hero h1 { margin-bottom: .3rem; color: var(--text); }
      .hero p { color: var(--muted); font-size: 0.95rem; }
      .badge { display:inline-block; padding:6px 10px; border-radius:999px; background:var(--glass); color:var(--accent); border:1px solid var(--border); font-size:0.8rem; }
      .cta-btn button { width:100%; background:linear-gradient(120deg,var(--accent),var(--accent-2)); color:#041021; border:none; }
      .upload-card, .chat-card {
        background: var(--card);
        border: 1px solid var(--border);
        border-radius: 16px;
        padding: 1rem 1.1rem;
        box-shadow: 0 12px 40px rgba(0,0,0,0.25);
      }
      .stExpander, .stFileUploader { color: var(--text); }
      .stChatMessage { background: transparent; }
      .stMarkdown p { color: var(--text); }
      .chat-card .stChatMessage[data-testid="stChatMessage"] div { color: var(--text); }
      .bike-loader { display:flex; gap:10px; font-size:30px; margin: 6px 0 2px; }
      .bike-loader div { animation: ride 0.9s ease-in-out infinite; }
      .bike-loader div:nth-child(2) { animation-delay: .15s; }
      .bike-loader div:nth-child(3) { animation-delay: .3s; }
      @keyframes ride { 0% { transform: translateX(0px); } 50% { transform: translateX(12px); } 100% { transform: translateX(0px); } }
    </style>
    """,
    unsafe_allow_html=True,
)

# --- 1. LOGO ---------------------------------------------------------------
LOGO_PATH = "1.png"  # pas eventueel aan als je logo-bestand anders heet

# --- 2. CONFIGURATIE & API ------------------------------------------------
try:
    if "GEMINI_API_KEY" in st.secrets:
        api_key = st.secrets["GEMINI_API_KEY"].strip()
        genai.configure(api_key=api_key)
    else:
        st.error("Geen API Key gevonden.")
        st.stop()
except Exception as e:
    st.error(f"Error: {e}")
    st.stop()

# --- 3. KENNIS LADEN (PDF & DOCX) -----------------------------------------
@st.cache_resource(show_spinner=False)
def load_all_knowledge():
    """Zoekt automatisch naar alle PDF en DOCX bestanden en leest ze."""
    combined_text = ""
    for filename in os.listdir("."):
        try:
            if filename.lower().endswith(".pdf"):
                reader = pypdf.PdfReader(filename)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        combined_text += text + "\n"
            elif filename.lower().endswith(".docx"):
                doc = docx.Document(filename)
                for para in doc.paragraphs:
                    combined_text += para.text + "\n"
        except Exception as e:
            print(f"Kon bestand {filename} niet lezen: {e}")
    return combined_text

knowledge_base = load_all_knowledge()

# --- 4. AI INSTRUCTIES ----------------------------------------------------
SYSTEM_PROMPT = f"""
ROL: Je bent een expert sportfysioloog van SportMetrics.

BRONMATERIAAL:
Je hebt toegang tot specifieke literatuur over trainingsleer (zie hieronder).
Gebruik DEZE INFORMATIE als de absolute waarheid.

=== START LITERATUUR ===
{knowledge_base}
=== EINDE LITERATUUR ===

BELANGRIJKE REGELS:
1. SportMetrics doet GEEN lactaatmetingen (prikken), alleen ademgasanalyse.
2. Gebruik de principes (zoals Seiler zones) zoals beschreven in de geüploade literatuur.
3. Wees praktisch, enthousiast en gebruik bulletpoints.
4. Geen medisch advies.
5. Geef altijd een props aan de persoon voor de test en bedankt dat hij of zij dit bij SportMetrics heeft gedaan.
"""

try:
    model = genai.GenerativeModel(
        model_name="gemini-2.5-flash",
        system_instruction=SYSTEM_PROMPT,
    )
except Exception as e:
    st.error(f"Model fout: {e}")

# --- 5. HERO --------------------------------------------------------------
hero = st.container()
with hero:
    col1, col2 = st.columns([1.7, 1.1])
    with col1:
        st.markdown("<span class='badge'>SportMetrics AI Coach</span>", unsafe_allow_html=True)
        st.markdown("<h1>🚴‍♂️ Jouw Wieler & Hardloop Expert</h1>", unsafe_allow_html=True)
        st.markdown("""
        Geef direct trainingsadvies op basis van je eigen rapporten en de beste literatuur.
        Upload je testresultaten of stel je vraag, wij vertalen het naar heldere zones en acties.
        """)
        st.write("‣ Praktisch en to-the-point · ‣ Seiler zones · ‣ Geen medisch advies · ‣ Altijd props voor jouw effort")
    with col2:
        if os.path.exists(LOGO_PATH):
            st.image(LOGO_PATH, width=260)
        else:
            st.info("Upload je logo als '1.png' in dezelfde map om het hier te tonen.")

# --- 6. UPLOAD ------------------------------------------------------------
upload_card = st.container()
with upload_card:
    st.markdown("<div class='upload-card'>", unsafe_allow_html=True)
    col_u1, col_u2 = st.columns([1.4, 1])
    with col_u1:
        st.subheader("📄 Upload je rapport")
        st.caption("PDF of DOCX, alles blijft lokaal.")
        uploaded_file = st.file_uploader("Kies je testresultaten", type=["pdf", "docx"], label_visibility="collapsed")
    with col_u2:
        st.markdown("<div class='cta-btn'>", unsafe_allow_html=True)
        st.button("🚀 Start analyse", use_container_width=True)
        st.markdown("</div>", unsafe_allow_html=True)
        st.caption("Tip: vraag 'Maak mijn zones' voor een kort overzicht.")
    st.markdown("</div>", unsafe_allow_html=True)

if uploaded_file is not None:
    try:
        client_pdf_text = ""
        if uploaded_file.name.lower().endswith(".pdf"):
            reader = pypdf.PdfReader(uploaded_file)
            for page in reader.pages:
                client_pdf_text += page.extract_text() + "\n"
        elif uploaded_file.name.lower().endswith(".docx"):
            doc = docx.Document(uploaded_file)
            for para in doc.paragraphs:
                client_pdf_text += para.text + "\n"
        st.session_state["last_uploaded_text"] = client_pdf_text
        st.toast("Rapport ontvangen! Typ je vraag beneden.", icon="✅")
    except Exception as e:
        st.error(f"Fout bij lezen rapport: {e}")

# --- 7. CHAT HISTORY ------------------------------------------------------
if "messages" not in st.session_state:
    st.session_state.messages = []
    intro = (
        "Hoi! Ik geef antwoord op basis van mijn AI-kennis en de best beschikbare literatuur over trainingsleer.\n\n"
        "Upload je testresultaten of stel direct een vraag!"
    )
    st.session_state.messages.append({"role": "assistant", "content": intro})

chat_box = st.container()
with chat_box:
    st.markdown("<div class='chat-card'>", unsafe_allow_html=True)
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    st.markdown("</div>", unsafe_allow_html=True)

# --- 8. CHAT INPUT --------------------------------------------------------
prompt = st.chat_input("Stel je vraag of zeg 'Maak mijn zones'...")

if prompt:
    extra_context = ""
    if "last_uploaded_text" in st.session_state:
        extra_context = f"\n\nHIER IS HET RAPPORT VAN DE KLANT:\n{st.session_state['last_uploaded_text']}\n\n"
        del st.session_state["last_uploaded_text"]

    full_prompt_for_ai = prompt + extra_context

    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    try:
        with st.chat_message("assistant"):
            bike_placeholder = st.empty()
            bike_placeholder.markdown(
                """
                <div class="bike-loader">
                  <div>🚴‍♀️</div><div>🚴‍♂️</div><div>🚴‍♀️</div>
                </div>
                <p style="color:var(--muted); margin-top:2px;">Antwoord wordt geladen...</p>
                """,
                unsafe_allow_html=True,
            )
            with st.spinner("🚴‍♂️🚴‍♀️ bezig met jouw advies..."):
                response = model.generate_content(full_prompt_for_ai)
            bike_placeholder.markdown(response.text)
            st.session_state.messages.append({"role": "assistant", "content": response.text})
    except Exception as e:
        st.error(f"De AI reageert niet: {e}")
